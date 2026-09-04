"""Importazione sicura di documenti policy in una bozza strutturata.

L'estrazione locale serve alla sandbox portfolio: rende espliciti i campi da
confermare senza modificare la policy attiva. PDF e DOCX vengono letti solo
quando le relative dipendenze sono installate.
"""

from __future__ import annotations

import ipaddress
import re
import socket
from html.parser import HTMLParser
from io import BytesIO
from pathlib import Path
from urllib.parse import urlparse

import requests

MAX_DOCUMENT_BYTES = 2_000_000
ALLOWED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx"}


class PolicyImportError(ValueError):
    """Errore mostrabile all'utente durante l'importazione."""


class _HTMLTextExtractor(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.parts: list[str] = []

    def handle_data(self, data: str) -> None:
        value = data.strip()
        if value:
            self.parts.append(value)


def _plain_html(value: str) -> str:
    parser = _HTMLTextExtractor()
    parser.feed(value)
    return "\n".join(parser.parts)


def document_text(filename: str, payload: bytes, content_type: str = "") -> str:
    if not payload:
        raise PolicyImportError("Il documento è vuoto.")
    if len(payload) > MAX_DOCUMENT_BYTES:
        raise PolicyImportError("Il documento supera il limite di 2 MB.")

    suffix = Path(filename or "policy.txt").suffix.lower()
    if suffix not in ALLOWED_EXTENSIONS and "html" not in content_type:
        raise PolicyImportError("Formato non supportato. Usa PDF, DOCX, TXT o MD.")

    try:
        if suffix == ".pdf" or content_type == "application/pdf":
            from pypdf import PdfReader

            reader = PdfReader(BytesIO(payload))
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
        elif suffix == ".docx" or "wordprocessingml" in content_type:
            from docx import Document

            document = Document(BytesIO(payload))
            text = "\n".join(paragraph.text for paragraph in document.paragraphs)
            for table in document.tables:
                for row in table.rows:
                    text += "\n" + " | ".join(cell.text for cell in row.cells)
        else:
            text = payload.decode("utf-8-sig", errors="replace")
            if "html" in content_type:
                text = _plain_html(text)
    except PolicyImportError:
        raise
    except Exception as exc:
        raise PolicyImportError("Non è stato possibile leggere il documento.") from exc

    if len(text.strip()) < 40:
        raise PolicyImportError("Il documento non contiene abbastanza testo analizzabile.")
    return text.strip()


def _validate_public_url(value: str) -> str:
    parsed = urlparse(value.strip())
    if parsed.scheme != "https" or not parsed.hostname:
        raise PolicyImportError("Inserisci un URL HTTPS pubblico.")
    try:
        addresses = {
            item[4][0]
            for item in socket.getaddrinfo(parsed.hostname, parsed.port or 443)
        }
    except socket.gaierror as exc:
        raise PolicyImportError("Il dominio indicato non è raggiungibile.") from exc
    for address in addresses:
        ip = ipaddress.ip_address(address)
        if any(
            (
                ip.is_private,
                ip.is_loopback,
                ip.is_link_local,
                ip.is_multicast,
                ip.is_reserved,
                ip.is_unspecified,
            )
        ):
            raise PolicyImportError("L'URL deve puntare a una risorsa pubblica.")
    return value.strip()


def url_text(value: str) -> tuple[str, str]:
    url = _validate_public_url(value)
    try:
        response = requests.get(
            url,
            allow_redirects=False,
            stream=True,
            timeout=(3.05, 10),
            headers={"User-Agent": "OpsPlaybookImporter/1.0"},
        )
        response.raise_for_status()
    except requests.RequestException as exc:
        raise PolicyImportError("Non è stato possibile scaricare la policy.") from exc
    if 300 <= response.status_code < 400:
        raise PolicyImportError("I redirect non sono supportati: usa l'URL finale.")

    try:
        declared_size = int(response.headers.get("content-length") or 0)
    except ValueError:
        declared_size = 0
    if declared_size > MAX_DOCUMENT_BYTES:
        raise PolicyImportError("Il documento supera il limite di 2 MB.")
    chunks: list[bytes] = []
    size = 0
    for chunk in response.iter_content(65_536):
        size += len(chunk)
        if size > MAX_DOCUMENT_BYTES:
            raise PolicyImportError("Il documento supera il limite di 2 MB.")
        chunks.append(chunk)

    content_type = response.headers.get("content-type", "").split(";", 1)[0]
    filename = Path(urlparse(url).path).name or (
        "policy.html" if "html" in content_type else "policy.txt"
    )
    return document_text(filename, b"".join(chunks), content_type), url


def _contains(text: str, *terms: str) -> bool:
    return any(term in text for term in terms)


def _generic_playbook_extraction(cleaned: str, lowered: str) -> dict:
    """Ordina procedure non e-commerce in una struttura minima verificabile."""
    agency_score = sum(term in lowered for term in ("brief", "scope", "deliverable", "asset", "progetto", "cliente approva"))
    internal_score = sum(term in lowered for term in ("accesso", "acquisto", "licenza", "incidente", "deroga", "richiesta interna"))
    kind = "agency" if agency_score > internal_score else "internal" if internal_score else "operations"

    def detected(terms: tuple[str, ...], positive: str) -> str:
        return positive if _contains(lowered, *terms) else "Da confermare"

    if kind == "agency":
        values = [
            ("intake_scope", "Scope minimo", detected(("scope", "obiettivo", "deliverable", "brief"), "Obiettivo e deliverable espliciti")),
            ("intake_deadline", "Scadenza", detected(("scadenza", "deadline", "consegna"), "Da verificare prima della pianificazione")),
            ("budget_gate", "Copertura economica", detected(("budget", "preventivo", "costo"), "Approvazione prima del kickoff")),
            ("change_control", "Cambio di scope", detected(("modifica", "change", "extra", "fuori scope"), "Valutare impatto su tempi e costi")),
            ("delivery_owner", "Responsabile", detected(("owner", "responsabile", "project manager"), "Owner obbligatorio")),
            ("approval_gate", "Approvazione finale", detected(("approv", "via libera", "conferma"), "Conferma umana esplicita")),
        ]
        sections = [
            ("Intake del lavoro", values[:3]),
            ("Delivery e controllo", values[3:]),
        ]
    elif kind == "internal":
        values = [
            ("request_reason", "Motivazione", detected(("motiv", "necessità", "obiettivo"), "Motivazione operativa obbligatoria")),
            ("request_owner", "Responsabile", detected(("owner", "responsabile", "assegn"), "Owner identificato")),
            ("approval_gate", "Approvazione", detected(("approv", "manager", "responsabile"), "Approvazione prima dell’esecuzione")),
            ("priority_rule", "Priorità", detected(("urgente", "priorità", "critico", "impatto"), "Definita da urgenza e impatto")),
            ("exception_rule", "Eccezioni", detected(("eccezione", "deroga", "fuori processo"), "Motivazione e scadenza obbligatorie")),
            ("audit_rule", "Tracciabilità", detected(("registra", "audit", "storico", "document"), "Decisione ed esito conservati")),
        ]
        sections = [
            ("Intake e responsabilità", values[:3]),
            ("Priorità, eccezioni e controllo", values[3:]),
        ]
    else:
        values = [
            ("request_type", "Tipo di richiesta", "Da confermare"),
            ("required_context", "Informazioni necessarie", "Da confermare"),
            ("decision_rule", "Criterio decisionale", "Da confermare"),
            ("owner_rule", "Responsabile", detected(("owner", "responsabile", "operatore"), "Responsabile esplicito")),
            ("escalation_rule", "Escalation", detected(("escal", "eccezione", "superiore"), "Percorso manuale previsto")),
            ("human_gate", "Controllo finale", detected(("approv", "revisione", "persona"), "Approvazione umana")),
        ]
        sections = [("Intake", values[:3]), ("Governance", values[3:])]

    rules = [
        {"id": rule_id, "label": label, "value": value, "confidence": 0.9 if value != "Da confermare" else 0.55, "needs_confirmation": value == "Da confermare"}
        for rule_id, label, value in values
    ]
    confirmations = []
    for line in cleaned.splitlines():
        if "DA CONFERMARE" in line.upper():
            description = re.sub(r"[`*#\[\]]", "", line).strip(" -:—")
            if description and description not in confirmations:
                confirmations.append(description)
    confirmations.extend(
        f"{item['label']}: il testo non definisce un valore univoco."
        for item in rules
        if item["needs_confirmation"]
    )
    normalized_document = [
        {"title": title, "items": [f"{label}: {value}" for _, label, value in section_rules]}
        for title, section_rules in sections
    ]
    return {
        "rules": rules,
        "normalized_document": normalized_document,
        "confirmations": confirmations[:7],
        "rule_count": len(rules),
        "confirmation_count": min(len(confirmations), 7),
        "engine": "structured_playbook_preview",
        "playbook_type": kind,
        "characters_read": len(cleaned),
    }


def extract_structured_rules(text: str) -> dict:
    """Converte il testo in una bozza esplicita da revisionare manualmente."""

    cleaned = text.strip()
    if len(cleaned) < 40:
        raise PolicyImportError("Inserisci almeno 40 caratteri di policy.")
    lowered = cleaned.casefold()
    if not _contains(lowered, "reso", "recesso", "garanzia", "rimborso", "doa"):
        return _generic_playbook_extraction(cleaned, lowered)
    day_match = re.search(r"\b(\d{1,3})\s*(?:giorni|giorno|gg)\b", lowered)
    window = f"{day_match.group(1)} giorni" if day_match else "Da confermare"

    starts_from = (
        "Data di consegna / tracking"
        if _contains(lowered, "data di consegna", "dalla consegna", "tracking")
        else "Da confermare"
    )
    condition = (
        "Integro e rivendibile"
        if _contains(lowered, "rivendibile", "stesse condizioni", "integro")
        else "Da confermare"
    )
    if "a carico del cliente" in lowered and "a carico azienda" in lowered:
        shipping = "Cliente per recesso · Azienda per difetti"
    elif "a carico del cliente" in lowered:
        shipping = "A carico del cliente"
    elif _contains(lowered, "a carico azienda", "a carico dell'azienda"):
        shipping = "A carico dell’azienda"
    else:
        shipping = "Da confermare"

    hygiene = (
        "Aperti: non idonei · Sigillati: idonei"
        if _contains(lowered, "rasoi", "rasoio", "spazzolini", "spazzolino")
        and _contains(lowered, "aperta", "aperto", "sigillo")
        else "Non specificato"
    )
    evidence = (
        "Foto e video richiesti"
        if "foto" in lowered and "video" in lowered
        else "Da confermare"
    )
    refund = (
        "Solo dopo il controllo fisico"
        if "rimborso" in lowered
        and _contains(lowered, "dopo che", "dopo il controllo", "verificat")
        else "Da confermare"
    )
    approval = (
        "Approvazione operatore obbligatoria"
        if "operatore" in lowered
        and _contains(lowered, "approva", "non invia mai", "revisione umana")
        else "Da confermare"
    )
    warranty_years = re.search(
        r"(?:garanzia|difett\w*)[^.\n]{0,80}?\b(\d{1,2})\s*ann[oi]\b", lowered
    )
    warranty_days = re.search(
        r"(?:garanzia|difett\w*)[^.\n]{0,80}?\b(\d{2,4})\s*(?:giorni|gg)\b", lowered
    )
    if warranty_years:
        years = int(warranty_years.group(1))
        warranty = f"{years} anni ({years * 365} giorni)"
    elif warranty_days:
        warranty = f"{warranty_days.group(1)} giorni"
    else:
        warranty = "Da confermare"

    values = [
        ("return_window", "Finestra recesso", window),
        ("starts_from", "Decorrenza", starts_from),
        ("product_condition", "Condizione prodotto", condition),
        ("return_shipping", "Spedizione di ritorno", shipping),
        ("hygiene_products", "Prodotti igienici aperti", hygiene),
        ("doa_evidence", "Prove DOA", evidence),
        ("warranty_window", "Finestra garanzia", warranty),
        ("refund_timing", "Avvio rimborso", refund),
        ("human_gate", "Approvazione finale", approval),
    ]
    rules = [
        {
            "id": rule_id,
            "label": label,
            "value": value,
            "confidence": 0.94 if value not in {"Da confermare", "Non specificato"} else 0.58,
            "needs_confirmation": value in {"Da confermare", "Non specificato"},
        }
        for rule_id, label, value in values
    ]

    confirmations = []
    for line in cleaned.splitlines():
        if "DA CONFERMARE" not in line.upper():
            continue
        description = re.sub(r"[`*#\[\]]", "", line).strip(" -:—")
        if description and description not in confirmations:
            confirmations.append(description)
    if not confirmations:
        confirmations = [
            f"{item['label']}: il documento non definisce un valore univoco."
            for item in rules
            if item["needs_confirmation"]
        ]

    by_id = {item["id"]: item["value"] for item in rules}
    normalized_document = [
        {
            "title": "Recesso",
            "items": [
                f"Finestra: {by_id['return_window']}",
                f"Decorrenza: {by_id['starts_from']}",
                f"Condizione richiesta: {by_id['product_condition']}",
                f"Costo del rientro: {by_id['return_shipping']}",
            ],
        },
        {
            "title": "Garanzia e difetti",
            "items": [
                f"Durata garanzia: {by_id['warranty_window']}",
                f"Prove richieste: {by_id['doa_evidence']}",
                f"Prodotti igienici: {by_id['hygiene_products']}",
            ],
        },
        {
            "title": "Controlli e responsabilità",
            "items": [
                f"Rimborso: {by_id['refund_timing']}",
                f"Approvazione: {by_id['human_gate']}",
                "Le regole ambigue restano sospese fino alla conferma umana.",
            ],
        },
    ]

    return {
        "rules": rules,
        "normalized_document": normalized_document,
        "confirmations": confirmations[:7],
        "rule_count": len(rules),
        "confirmation_count": min(len(confirmations), 7),
        "engine": "structured_extraction_preview",
        "characters_read": len(cleaned),
    }
